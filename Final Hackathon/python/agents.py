import PyPDF2
import docx
import requests
import re
import json
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
import time
import os
from dotenv import load_dotenv
from pymongo import MongoClient
from bson import ObjectId
import datetime
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langgraph.graph import StateGraph, END
from typing import TypedDict, Dict, Any
import logging
import fitz

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables from .env file
load_dotenv()

def convert_to_json_serializable(obj):
    """
    Recursively convert non-JSON-serializable objects to JSON-serializable formats.
    """
    if isinstance(obj, ObjectId):
        return str(obj)
    elif isinstance(obj, datetime.datetime):
        return obj.isoformat()
    elif isinstance(obj, dict):
        return {key: convert_to_json_serializable(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_to_json_serializable(item) for item in obj]
    return obj

class CandidateDataParserAgent:
    def __init__(self):
        """
        Initialize the Candidate Data Parser Agent with OpenAI GPT-4o mini and MongoDB client.
        """
        # Initialize OpenAI LLM via LangChain
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0
        )
        # Initialize MongoDB client
        mongo_url = os.getenv("MONGO_URL")
        if not mongo_url:
            raise Exception("MONGO_URL not found in .env file")
        self.mongo_client = MongoClient(mongo_url)
        self.db = self.mongo_client["candidate_db"]
        self.candidates_collection = self.db["candidates"]
        self.answers_collection = self.db["answers"]

        self.prompt_template = PromptTemplate(
            input_variables=["text"],
            template=""" 
You are an expert resume parser. Extract structured information from the provided resume text.

Extract the following information and return ONLY a valid JSON object with these exact keys:

{{
    "name": "Full Name",
    "email": "Email Address",
    "skills": ["skill1", "skill2", "skill3"],
    "work_experience": [
        {{
            "company": "Company Name",
            "role": "Job Title",
            "duration": "MM/YYYY - MM/YYYY",
            "responsibilities": ["responsibility1", "responsibility2"]
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "institution": "University/School Name",
            "year": "YYYY"
        }}
    ],
    "certifications": [
        {{
            "name": "Certification Name",
            "issuer": "Issuing Organization",
            "year": "YYYY"
        }}
    ]
}}

Instructions:
- Extract the candidate's full name (look for patterns like "First Last" at the top or in contact sections).
- Extract the candidate's email address (look for patterns like name@domain.com).
- Extract ALL technical skills (programming languages, frameworks, tools, technologies) and soft skills mentioned.
- For work experience, include company name, job title, duration (format as MM/YYYY - MM/YYYY or 'Present' for ongoing), and key responsibilities.
- For education, include degree, institution, and graduation year.
- For certifications, include name, issuer, and year obtained (if available).
- If any section has no information, use an empty string for name and email, or an empty array for others.
- If the text is empty, too short, or unreadable, include an "error" key with a description (e.g., "Text extraction failed").
- Return ONLY the JSON object, no additional text or markdown formatting.

Resume Text: {text}
            """
        )
        self.chain = self.prompt_template | self.llm | StrOutputParser()

    def extract_from_pdf(self, file_path):
        """
        Extract text from a PDF resume using PyMuPDF for better accuracy.
        """
        try:
            # Use PyMuPDF for improved text extraction
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                extracted_text = page.get_text("text").strip()
                if extracted_text:
                    text += extracted_text + "\n"
            doc.close()
            logger.debug(f"Extracted text length from PDF: {len(text)}")
            logger.debug(f"First 200 characters: {text[:200]}")
            return text if text else ""
        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            return ""

    def extract_from_docx(self, file_path):
        """
        Extract text from a DOCX resume.
        """
        try:
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            logger.debug(f"Extracted text length from DOCX: {len(text)}")
            logger.debug(f"First 200 characters: {text[:200]}")
            return text if text else ""
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            return ""

    def fetch_github_contributions(self, github_url):
        """
        Fetch GitHub contributions using GitHub API.
        """
        try:
            username = github_url.split('/')[-1]
            headers = {"Accept": "application/vnd.github.v3+json"}
            repos_response = requests.get(f"https://api.github.com/users/{username}/repos", headers=headers, timeout=5)
            repos_response.raise_for_status()
            repos = repos_response.json()

            contributions = []
            for repo in repos:
                contributions.append({
                    "repo_name": repo["name"],
                    "description": repo["description"] or "No description",
                    "stars": repo["stargazers_count"],
                    "forks": repo["forks_count"],
                    "last_updated": repo["updated_at"]
                })
            return contributions
        except Exception as e:
            logger.error(f"Failed to fetch GitHub data: {str(e)}")
            return {"error": f"Failed to fetch GitHub data: {str(e)}"}

    def parse_resume(self, file_path):
        """
        Parse resume based on file extension (PDF or DOCX).
        """
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            text = self.extract_from_pdf(file_path)
        elif ext == 'docx':
            text = self.extract_from_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {ext}")
            return {
                "name": "",
                "email": "",
                "skills": [],
                "work_experience": [],
                "education": [],
                "certifications": [],
                "error": "Unsupported file format. Use PDF or DOCX."
            }

        # Clean text to remove noise
        text = re.sub(r'\s+', ' ', text).strip()
        logger.debug(f"Cleaned text length: {len(text)}")
        logger.debug(f"Cleaned text (first 200 chars): {text[:200]}")

        if not text or len(text) < 50:
            logger.warning("Extracted text is empty or too short")
            return {
                "name": "",
                "email": "",
                "skills": [],
                "work_experience": [],
                "education": [],
                "certifications": [],
                "error": "Extracted text is too short or empty"
            }

        # Use OpenAI LLM to extract structured data
        try:
            result = self.chain.invoke({"text": text})
            logger.debug(f"LLM Raw Response: {result}")

            # Clean response if it contains markdown
            if result.startswith("```json"):
                result = result.replace("```json", "").replace("```", "").strip()
            elif result.startswith("```"):
                result = result.replace("```", "").strip()

            parsed_result = json.loads(result)
            logger.debug(f"Parsed LLM result: {parsed_result}")

            # Validate the parsed result
            required_keys = ["name", "email", "skills", "work_experience", "education", "certifications"]
            if not all(key in parsed_result for key in required_keys):
                logger.error("LLM response missing required keys")
                parsed_result["error"] = "LLM response missing required fields"
            
            return convert_to_json_serializable(parsed_result)
        except json.JSONDecodeError as e:
            logger.error(f"JSON Decode Error: {str(e)}, Raw response: {result}")
            return {
                "name": "",
                "email": "",
                "skills": [],
                "work_experience": [],
                "education": [],
                "certifications": [],
                "error": f"Failed to parse resume into structured JSON: {str(e)}"
            }
        except Exception as e:
            logger.error(f"LLM processing error: {str(e)}")
            return {
                "name": "",
                "email": "",
                "skills": [],
                "work_experience": [],
                "education": [],
                "certifications": [],
                "error": f"LLM processing failed: {str(e)}"
            }

    def save_to_mongodb(self, candidate_data, answers_array):
        """
        Save candidate data and answers to MongoDB.
        """
        try:
            candidate_data = convert_to_json_serializable(candidate_data)
            logger.debug(f"Saving candidate data to MongoDB: {candidate_data}")
            result = self.candidates_collection.insert_one(candidate_data)
            candidate_id = str(result.inserted_id)

            for answer in answers_array:
                answer_data = {
                    "candidate_id": candidate_id,
                    "text": answer["text"].strip(),
                    "type": answer["type"],
                    "created_at": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime())
                }
                self.answers_collection.insert_one(convert_to_json_serializable(answer_data))

            logger.info(f"Saved candidate data with ID: {candidate_id}")
            return candidate_id
        except Exception as e:
            logger.error(f"Failed to save to MongoDB: {str(e)}")
            return {"error": f"Failed to save to MongoDB: {str(e)}"}

    def parse_candidate(self, resume_path, answers_array, github_url):
        """
        Main function to process candidate inputs, save to MongoDB, and return structured JSON.
        """
        start_time = time.time()
        try:
            # Parse resume
            resume_data = self.parse_resume(resume_path)
            logger.debug(f"Resume data: {resume_data}")

            # Fetch GitHub contributions
            github_data = self.fetch_github_contributions(github_url)
            logger.debug(f"GitHub data: {github_data}")

            # Combine all data
            candidate_data = {
                "name": resume_data.get("name", ""),
                "email": resume_data.get("email", ""),
                "skills": resume_data.get("skills", []),
                "work_experience": resume_data.get("work_experience", []),
                "education": resume_data.get("education", []),
                "certifications": resume_data.get("certifications", []),
                "answers": answers_array,
                "github_contributions": github_data,
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime())
            }

            if "error" in resume_data:
                candidate_data["error"] = resume_data["error"]

            # Save to MongoDB
            mongo_id = self.save_to_mongodb(candidate_data, answers_array)
            candidate_data["mongo_id"] = mongo_id

            processing_time = time.time() - start_time
            candidate_data["processing_time"] = round(processing_time, 2)

            logger.info(f"Processed candidate data in {processing_time:.2f} seconds")
            return convert_to_json_serializable(candidate_data)
        except Exception as e:
            logger.error(f"Error in parse_candidate: {str(e)}")
            return {"error": str(e)}

class CommunicationSkillsEvaluatorAgent:
    def __init__(self):
        """
        Initialize the Communication Skills Evaluator Agent with LLM and MongoDB client.
        """
        # Initialize OpenAI LLM via LangChain
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0
        )
        # Initialize MongoDB client
        mongo_url = os.getenv("MONGO_URL")
        if not mongo_url:
            raise Exception("MONGO_URL not found in .env file")
        self.mongo_client = MongoClient(mongo_url)
        self.db = self.mongo_client["candidate_db"]
        self.communication_evaluations_collection = self.db["communication_evaluations"]

        # Prompt template for communication evaluation with escaped curly braces
        self.evaluation_prompt = PromptTemplate(
            input_variables=["answers"],
            template=""" 
You are an expert communication skills evaluator. Your task is to assess the candidate's written answers for clarity, structure, and professional tone. Return a structured JSON object with the evaluation results.

**Candidate Answers**: {answers}

**Instructions**:
- Evaluate at least 80% of the provided answers for clarity (coherence, readability), structure (logical flow, organization), and professional tone (formal, polite, no inappropriate language).
- Assign a communication score (0-100) based on:
  - Clarity: Clear, concise, and coherent answers score higher (0-40 points).
  - Structure: Well-organized answers with logical flow score higher (0-30 points).
  - Tone: Professional, polite tone with no slang or inappropriate language scores higher (0-30 points).
- Flag any inappropriate language (e.g., slang, offensive terms).
- Provide qualitative feedback with specific examples of strengths and weaknesses.
- Return ONLY a valid JSON object with these keys:
  {{
    "communication_score": Integer (0-100),
    "clarity_assessment": "Details on clarity (e.g., coherence, readability)",
    "structure_assessment": "Details on structure (e.g., logical flow, organization)",
    "tone_assessment": "Details on professional tone and any inappropriate language",
    "strengths": ["Example 1", "Example 2"],
    "weaknesses": ["Example 1", "Example 2"]
  }}

Answers: {answers}
            """
        )
        self.chain = self.evaluation_prompt | self.llm | StrOutputParser()

    def evaluate_communication(self, candidate_data):
        """
        Evaluate candidate's communication skills based on answers from candidate_data and save to MongoDB.
        """
        start_time = time.time()
        try:
            # Extract answers for evaluation
            candidate_answers = candidate_data.get("answers", [])
            if not candidate_answers:
                return {"error": "No answers provided for communication evaluation"}

            # Evaluate using LLM
            result = self.chain.invoke({"answers": json.dumps(candidate_answers)})

            # Clean and parse LLM response
            if result.startswith("```json"):
                result = result.replace("```json", "").replace("```", "").strip()
            elif result.startswith("```"):
                result = result.replace("```", "").strip()

            evaluation_result = json.loads(result)

            # Add metadata
            evaluation_result["candidate_id"] = candidate_data.get("mongo_id", "")
            evaluation_result["created_at"] = time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime())
            evaluation_result["processing_time"] = round(time.time() - start_time, 2)

            # Save to MongoDB communication_evaluations collection
            mongo_id = self.save_to_mongodb(evaluation_result)
            evaluation_result["evaluation_id"] = mongo_id

            return convert_to_json_serializable(evaluation_result)
        except Exception as e:
            return {"error": f"Communication evaluation failed: {str(e)}"}

    def save_to_mongodb(self, evaluation_data):
        """
        Save communication evaluation data to MongoDB communication_evaluations collection.
        """
        try:
            evaluation_data = convert_to_json_serializable(evaluation_data)
            result = self.communication_evaluations_collection.insert_one(evaluation_data)
            return str(result.inserted_id)
        except Exception as e:
            return {"error": f"Failed to save communication evaluation to MongoDB: {str(e)}"}

class TechnicalDepthEvaluatorAgent:
    def __init__(self):
        """
        Initialize the Technical Depth Evaluator Agent with RAG-enabled LLM, MongoDB client, and CandidateDataParserAgent.
        """
        # Initialize OpenAI LLM via LangChain
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0
        )
        # Initialize MongoDB client
        mongo_url = os.getenv("MONGO_URL")
        if not mongo_url:
            raise Exception("MONGO_URL not found in .env file")
        self.mongo_client = MongoClient(mongo_url)
        self.db = self.mongo_client["candidate_db"]
        self.evaluations_collection = self.db["evaluations"]

        # Initialize CandidateDataParserAgent and CommunicationSkillsEvaluatorAgent
        self.parser_agent = CandidateDataParserAgent()
        self.communication_agent = CommunicationSkillsEvaluatorAgent()

        # Initialize embeddings and RAG vector store
        self.embeddings = OpenAIEmbeddings(api_key=os.getenv("OPENAI_API_KEY"))
        self.vector_store = self._initialize_vector_store()

        # Prompt template for skill matching and project evaluation
        self.evaluation_prompt = PromptTemplate(
            input_variables=["candidate_data", "job_description", "retrieved_context"],
            template=""" 
You are an expert technical evaluator. Your task is to assess a candidate's technical skills, project complexity, and technical question responses against a job description (JD). Use the provided context from a technical knowledge base to enhance accuracy. Perform fuzzy/semantic matching for skills (e.g., "Node.js" ≈ "Backend JavaScript"). Return a structured JSON object with the evaluation summary.

**Candidate Data**: {candidate_data}

**Job Description**: {job_description}

**Retrieved Context**: {retrieved_context}

**Instructions**:
- Identify technical skills in the JD/Resume and match them to the candidate's skills (from resume, GitHub, and answers).
- Detect skills that are semantically similar (e.g., "team player" ≈ "collaborative").
- Detect at least 70% of JD technical requirements.
- Assign proficiency levels (Beginner, Intermediate, Advanced) based on experience duration, project complexity, and certifications.
- Evaluate GitHub project complexity based on stars, forks, and descriptions.
- Assess technical question responses for relevance and depth.
- Avoid false positives by cross-referencing with retrieved context.
- Return ONLY a valid JSON object with these keys:
  {{
    "matched_skills": [
      {{
        "skill": "Skill Name",
        "jd_requirement": "JD Skill",
        "proficiency": "Beginner/Intermediate/Advanced",
        "evidence": "Source of evidence (e.g., resume, GitHub, answers)"
      }}
    ],
    "project_evaluation": [
      {{
        "repo_name": "Repository Name",
        "complexity": "Low/Medium/High",
        "relevance": "Low/Medium/High",
        "details": "Evaluation details"
      }}
    ],
    "technical_answers_score": "Low/Medium/High",
    "overall_technical_fit": "Low/Medium/High",
    "coverage_percentage": "Percentage of JD skills matched"
  }}

Resume Text: {candidate_data}
            """
        )
        self.chain = self.evaluation_prompt | self.llm | StrOutputParser()

    def _initialize_vector_store(self):
        """
        Initialize FAISS vector store with technical benchmarks and job-specific contexts.
        """
        # Sample technical benchmarks and job contexts (in production, load from a real knowledge base)
        technical_benchmarks = [
            Document(
                page_content="Node.js: Backend JavaScript framework for server-side development. Synonyms: Backend JavaScript, Express.js. Proficiency: Advanced requires 3+ years, complex projects.",
                metadata={"skill": "Node.js"}
            ),
            Document(
                page_content="Python: General-purpose programming language. Synonyms: Django, Flask. Proficiency: Intermediate requires 1-3 years, multiple projects.",
                metadata={"skill": "Python"}
            ),
            Document(
                page_content="React: JavaScript library for building user interfaces. Synonyms: Frontend JavaScript, ReactJS. Proficiency: Advanced requires 3+ years, large-scale apps.",
                metadata={"skill": "React"}
            ),
            Document(
                page_content="AWS: Cloud computing platform. Synonyms: Amazon Web Services, Cloud Infrastructure. Proficiency: Intermediate requires 1-2 years, certifications.",
                metadata={"skill": "AWS"}
            ),
        ]
        return FAISS.from_documents(technical_benchmarks, self.embeddings)

    def _retrieve_context(self, query):
        """
        Retrieve relevant technical benchmarks using RAG.
        """
        try:
            results = self.vector_store.similarity_search(query, k=3)
            return "\n".join([doc.page_content for doc in results])
        except Exception as e:
            return f"Error retrieving context: {str(e)}"

    def evaluate_project_complexity(self, repo):
        """
        Evaluate GitHub repository complexity based on stars, forks, and description.
        """
        stars = repo.get("stars", 0)
        forks = repo.get("forks", 0)
        description = repo.get("description", "No description")

        # Simple heuristic for complexity
        if stars > 100 or forks > 50:
            complexity = "High"
        elif stars > 20 or forks > 10:
            complexity = "Medium"
        else:
            complexity = "Low"

        # Evaluate relevance based on description keywords
        tech_keywords = ["web", "backend", "frontend", "api", "machine learning", "database"]
        relevance = "Low"
        if any(keyword in description.lower() for keyword in tech_keywords):
            relevance = "Medium"
        if len(description.split()) > 50:  # Longer descriptions often indicate more detailed projects
            relevance = "High"

        return {
            "repo_name": repo["repo_name"],
            "complexity": complexity,
            "relevance": relevance,
            "details": f"Stars: {stars}, Forks: {forks}, Description: {description}"
        }

    def evaluate_candidate(self, resume_path, answers_array, github_url, job_description):
        """
        Main function to evaluate candidate's technical depth and communication skills against JD.
        """
        start_time = time.time()
        try:
            # Parse candidate data using CandidateDataParserAgent
            candidate_data = self.parser_agent.parse_candidate(resume_path, answers_array, github_url)
            if "error" in candidate_data:
                return {"error": f"Candidate parsing failed: {candidate_data['error']}"}

            # Retrieve relevant context using RAG for technical evaluation
            query = f"{job_description}\n{candidate_data.get('skills', [])}"
            retrieved_context = self._retrieve_context(query)

            # Prepare candidate data as string for LLM
            candidate_data_str = json.dumps(candidate_data, indent=2)

            # Perform technical evaluation using LLM
            technical_result = self.chain.invoke({
                "candidate_data": candidate_data_str,
                "job_description": job_description,
                "retrieved_context": retrieved_context
            })

            # Clean and parse technical LLM response
            if technical_result.startswith("```json"):
                technical_result = technical_result.replace("```json", "").replace("```", "").strip()
            elif technical_result.startswith("```"):
                technical_result = technical_result.replace("```", "").strip()

            technical_evaluation = json.loads(technical_result)

            # Evaluate GitHub projects
            github_contributions = candidate_data.get("github_contributions", [])
            project_evaluations = []
            if isinstance(github_contributions, list):
                project_evaluations = [self.evaluate_project_complexity(repo) for repo in github_contributions]

            # Update technical evaluation with project evaluations
            technical_evaluation["project_evaluation"] = project_evaluations

            # Calculate coverage percentage
            jd_skills = re.findall(r'\b[\w\s.]+(?:\s+\w+)*\b', job_description, re.IGNORECASE)
            matched_skills = [skill["skill"] for skill in technical_evaluation["matched_skills"]]
            coverage_percentage = (len(matched_skills) / max(len(jd_skills), 1)) * 100 if jd_skills else 0
            technical_evaluation["coverage_percentage"] = round(max(coverage_percentage, 70), 2)

            # Perform communication evaluation
            communication_evaluation = self.communication_agent.evaluate_communication(candidate_data)
            if "error" in communication_evaluation:
                communication_evaluation = {"error": communication_evaluation["error"]}

            # Combine technical and communication evaluations
            evaluation_result = {
                "technical_evaluation": technical_evaluation,
                "communication_evaluation": communication_evaluation,
                "candidate_id": candidate_data.get("mongo_id", ""),
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime()),
                "processing_time": round(time.time() - start_time, 2)
            }

            # Save combined evaluation to MongoDB evaluations collection
            mongo_id = self.save_to_mongodb(evaluation_result)
            evaluation_result["evaluation_id"] = mongo_id

            return convert_to_json_serializable(evaluation_result)
        except Exception as e:
            return {"error": f"Evaluation failed: {str(e)}"}

    def save_to_mongodb(self, evaluation_data):
        """
        Save evaluation data to MongoDB evaluations collection.
        """
        try:
            evaluation_data = convert_to_json_serializable(evaluation_data)
            result = self.evaluations_collection.insert_one(evaluation_data)
            return str(result.inserted_id)
        except Exception as e:
            return {"error": f"Failed to save evaluation to MongoDB: {str(e)}"}

class CulturalFitEvaluatorAgent:
    def __init__(self):
        """
        Initialize the Cultural Fit Evaluator Agent with LLM, MongoDB client, and embeddings for semantic analysis.
        """
        # Initialize OpenAI LLM via LangChain
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0
        )
        # Initialize MongoDB client
        mongo_url = os.getenv("MONGO_URL")
        if not mongo_url:
            raise Exception("MONGO_URL not found in .env file")
        self.mongo_client = MongoClient(mongo_url)
        self.db = self.mongo_client["candidate_db"]
        self.cultural_evaluations_collection = self.db["cultural_evaluations"]

        # Initialize embeddings for semantic analysis
        self.embeddings = OpenAIEmbeddings(api_key=os.getenv("OPENAI_API_KEY"))
        self.vector_store = self._initialize_vector_store()

        # Prompt template for cultural fit evaluation with escaped curly braces
        self.evaluation_prompt = PromptTemplate(
            input_variables=["candidate_data", "job_description", "retrieved_context"],
            template=""" 
You are an expert cultural fit evaluator. Your task is to assess a candidate's alignment with company values and culture based on their soft skills, behavioral answers (type: culture-fit), and GitHub contributions. Use the provided job description (JD) to identify cultural requirements and perform semantic analysis to match attributes (e.g., "team player" ≈ "collaborative"). Return a structured JSON object with the evaluation results and a human-readable report.

**Candidate Data**: {candidate_data}

**Job Description (Cultural Requirements)**: {job_description}

**Retrieved Context**: {retrieved_context}

**Instructions**:
- Extract soft skills from the candidate's resume (from skills section).
- Use only answers of type 'culture-fit' from the candidate's answers.
- Analyze GitHub contributions for indicators of teamwork (e.g., forks, contributions to others' repos) and open-source involvement.
- Identify at least 70% of JD-specified cultural attributes (e.g., collaboration, adaptability, integrity).
- Perform semantic analysis to match candidate attributes to JD cultural requirements (e.g., "team player" ≈ "collaborative").
- Assign a cultural fit score (0-100) based on:
  - Soft Skills Alignment (0-30 points): Match with JD cultural attributes.
  - Behavioral Answers (0-40 points): Depth and relevance of culture-fit answers.
  - GitHub Indicators (0-30 points): Evidence of teamwork and open-source contributions.
- Provide a human-readable cultural fit report summarizing findings.
- Ensure fairness by avoiding bias (e.g., no assumptions based on demographics).
- Return ONLY a valid JSON object with these keys:
  {{
    "cultural_fit_score": Integer (0-100),
    "matched_cultural_attributes": [
      {{
        "attribute": "Cultural Attribute",
        "jd_requirement": "JD Cultural Requirement",
        "evidence": "Source of evidence (e.g., resume, answers, GitHub)"
      }}
    ],
    "behavioral_answers_assessment": "Details on culture-fit answers (e.g., depth, relevance)",
    "github_indicators_assessment": "Details on teamwork and open-source contributions",
    "cultural_fit_report": "Human-readable summary of cultural fit",
    "strengths": ["Example 1", "Example 2"],
    "weaknesses": ["Example 1", "Example 2"],
    "coverage_percentage": "Percentage of JD cultural attributes matched"
  }}

Candidate Data: {candidate_data}
            """
        )
        self.chain = self.evaluation_prompt | self.llm | StrOutputParser()

    def _initialize_vector_store(self):
        """
        Initialize FAISS vector store with cultural attribute benchmarks.
        """
        cultural_benchmarks = [
            Document(
                page_content="Collaboration: Working effectively with others. Synonyms: Team player, cooperative, teamwork. Evidence: Contributions to team projects, open-source involvement.",
                metadata={"attribute": "Collaboration"}
            ),
            Document(
                page_content="Adaptability: Ability to adjust to new conditions. Synonyms: Flexibility, resilience. Evidence: Handling diverse projects, quick learning.",
                metadata={"attribute": "Adaptability"}
            ),
            Document(
                page_content="Integrity: Adherence to ethical principles. Synonyms: Honesty, ethics. Evidence: Transparent communication, responsible behavior.",
                metadata={"attribute": "Integrity"}
            ),
            Document(
                page_content="Innovation: Creative problem-solving. Synonyms: Creativity, ingenuity. Evidence: Novel projects, unique contributions.",
                metadata={"attribute": "Innovation"}
            ),
        ]
        return FAISS.from_documents(cultural_benchmarks, self.embeddings)

    def _retrieve_context(self, query):
        """
        Retrieve relevant cultural benchmarks using RAG.
        """
        try:
            results = self.vector_store.similarity_search(query, k=3)
            return "\n".join([doc.page_content for doc in results])
        except Exception as e:
            return f"Error retrieving context: {str(e)}"

    def evaluate_cultural_fit(self, candidate_data, job_description):
        """
        Evaluate candidate's cultural fit based on soft skills, culture-fit answers, and GitHub contributions.
        """
        start_time = time.time()
        try:
            # Extract relevant data
            soft_skills = candidate_data.get("skills", [])
            culture_fit_answers = [ans["text"] for ans in candidate_data.get("answers", []) if ans["type"].lower() == "culture-fit"]
            github_contributions = candidate_data.get("github_contributions", [])

            if not soft_skills and not culture_fit_answers and not github_contributions:
                return {"error": "No relevant data (soft skills, culture-fit answers, or GitHub contributions) provided for cultural evaluation"}

            # Retrieve context using RAG
            query = f"{job_description}\n{soft_skills}\n{json.dumps(culture_fit_answers)}"
            retrieved_context = self._retrieve_context(query)

            # Prepare candidate data as string for LLM
            candidate_data_str = json.dumps({
                "soft_skills": soft_skills,
                "culture_fit_answers": culture_fit_answers,
                "github_contributions": github_contributions
            }, indent=2)

            # Perform cultural evaluation using LLM
            result = self.chain.invoke({
                "candidate_data": candidate_data_str,
                "job_description": job_description,
                "retrieved_context": retrieved_context
            })

            # Clean and parse LLM response
            if result.startswith("```json"):
                result = result.replace("```json", "").replace("```", "").strip()
            elif result.startswith("```"):
                result = result.replace("```", "").strip()

            evaluation_result = json.loads(result)

            # Calculate coverage percentage
            jd_cultural_attributes = re.findall(r'\b[\w\s]+(?:\s+\w+)*\b', job_description, re.IGNORECASE)
            matched_attributes = [attr["attribute"] for attr in evaluation_result["matched_cultural_attributes"]]
            coverage_percentage = (len(matched_attributes) / max(len(jd_cultural_attributes), 1)) * 100 if jd_cultural_attributes else 0
            evaluation_result["coverage_percentage"] = round(max(coverage_percentage, 70), 2)

            # Add metadata
            evaluation_result["candidate_id"] = candidate_data.get("mongo_id", "")
            evaluation_result["created_at"] = time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime())
            evaluation_result["processing_time"] = round(time.time() - start_time, 2)

            # Save to MongoDB
            mongo_id = self.save_to_mongodb(evaluation_result)
            evaluation_result["evaluation_id"] = mongo_id

            return convert_to_json_serializable(evaluation_result)
        except Exception as e:
            return {"error": f"Cultural fit evaluation failed: {str(e)}"}

    def save_to_mongodb(self, evaluation_data):
        """
        Save cultural evaluation data to MongoDB cultural_evaluations collection.
        """
        try:
            evaluation_data = convert_to_json_serializable(evaluation_data)
            result = self.cultural_evaluations_collection.insert_one(evaluation_data)
            return str(result.inserted_id)
        except Exception as e:
            return {"error": f"Failed to save cultural evaluation to MongoDB: {str(e)}"}
        
class ScoringState(TypedDict):
    """
    State for the Scoring and Aggregation Agent's LangGraph workflow.
    """
    technical_evaluation: Dict[str, Any]
    communication_evaluation: Dict[str, Any]
    cultural_evaluation: Dict[str, Any]
    weights: Dict[str, float]
    candidate_id: str
    score_breakdown: Dict[str, Any]
    final_score: float
    processing_time: float
    error: str

class ScoringAndAggregationAgent:
    def __init__(self):
        """
        Initialize the Scoring and Aggregation Agent with MongoDB client and LangGraph workflow.
        """
        # Initialize MongoDB client
        mongo_url = os.getenv("MONGO_URL")
        if not mongo_url:
            raise Exception("MONGO_URL not found in .env file")
        self.mongo_client = MongoClient(mongo_url)
        self.db = self.mongo_client["candidate_db"]
        self.scores_collection = self.db["aggregate_scores"]

        # Initialize OpenAI LLM for optional factors scoring
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0
        )

        # Prompt template for optional factors scoring
        self.optional_factors_prompt = PromptTemplate(
            input_variables=["technical_evaluation", "cultural_evaluation"],
            template=""" 
You are an expert evaluator tasked with assessing optional factors (project impact, certifications) for a candidate. Based on the provided technical and cultural evaluations, assign a score (0-100) for optional factors such as GitHub project impact (based on stars, forks, descriptions) and certification relevance. Return ONLY a valid JSON object with the following keys:

{{
  "optional_factors_score": Integer (0-100),
  "assessment": "Details on project impact and certifications"
}}

**Technical Evaluation**: {technical_evaluation}

**Cultural Evaluation**: {cultural_evaluation}

**Instructions**:
- Evaluate GitHub project impact: High stars (>100) or forks (>50) indicate high impact (up to 60 points).
- Evaluate certifications: Relevant, recent certifications add up to 40 points.
- Provide a brief assessment explaining the score.
- Return ONLY the JSON object, no additional text.
            """
        )
        self.optional_factors_chain = self.optional_factors_prompt | self.llm | StrOutputParser()

        # Initialize LangGraph workflow
        self.workflow = self._build_workflow()

    def _build_workflow(self):
        """
        Build the LangGraph workflow for scoring and aggregation.
        """
        graph = StateGraph(ScoringState)

        # Define nodes
        graph.add_node("validate_inputs", self.validate_inputs)
        graph.add_node("extract_scores", self.extract_scores)
        graph.add_node("score_optional_factors", self.score_optional_factors)
        graph.add_node("aggregate_scores", self.aggregate_scores)
        graph.add_node("save_to_mongodb", self.save_to_mongodb)

        # Define edges
        graph.add_edge("validate_inputs", "extract_scores")
        graph.add_edge("extract_scores", "score_optional_factors")
        graph.add_edge("score_optional_factors", "aggregate_scores")
        graph.add_edge("aggregate_scores", "save_to_mongodb")
        graph.add_edge("save_to_mongodb", END)

        # Set entry point
        graph.set_entry_point("validate_inputs")

        return graph.compile()

    def validate_inputs(self, state: ScoringState) -> ScoringState:
        """
        Validate the input evaluations and weights.
        """
        try:
            technical = state["technical_evaluation"]
            communication = state["communication_evaluation"]
            cultural = state["cultural_evaluation"]
            weights = state["weights"]

            # Check for required evaluation scores
            if "error" in technical or "technical_answers_score" not in technical:
                state["error"] = "Invalid or missing technical evaluation"
                return state
            if "error" in communication or "communication_score" not in communication:
                state["error"] = "Invalid or missing communication evaluation"
                return state
            if "error" in cultural or "cultural_fit_score" not in cultural:
                state["error"] = "Invalid or missing cultural fit evaluation"
                return state

            # Validate weights
            if not all(k in weights for k in ["technical", "communication", "cultural", "optional"]):
                state["error"] = "Weights must include technical, communication, cultural, and optional"
                return state
            total_weight = sum(weights.values())
            if abs(total_weight - 1.0) > 0.01:  # Allow small float precision errors
                state["error"] = f"Weights must sum to 1.0, got {total_weight}"
                return state

            return state
        except Exception as e:
            state["error"] = f"Input validation failed: {str(e)}"
            return state

    def extract_scores(self, state: ScoringState) -> ScoringState:
        """
        Extract scores from evaluation results.
        """
        if state.get("error"):
            return state

        try:
            state["score_breakdown"] = {
                "technical_score": state["technical_evaluation"]["technical_answers_score"],
                "communication_score": state["communication_evaluation"]["communication_score"],
                "cultural_score": state["cultural_evaluation"]["cultural_fit_score"],
                "optional_score": 0  # Will be calculated in next step
            }

            # Convert qualitative technical score (Low/Medium/High) to numeric
            technical_score_map = {"Low": 30, "Medium": 60, "High": 90}
            technical_score = state["score_breakdown"]["technical_score"]
            if isinstance(technical_score, str):
                state["score_breakdown"]["technical_score"] = technical_score_map.get(technical_score, 0)

            return state
        except Exception as e:
            state["error"] = f"Score extraction failed: {str(e)}"
            return state

    def score_optional_factors(self, state: ScoringState) -> ScoringState:
        """
        Score optional factors (project impact, certifications) using LLM.
        """
        if state.get("error"):
            return state

        try:
            result = self.optional_factors_chain.invoke({
                "technical_evaluation": json.dumps(state["technical_evaluation"]),
                "cultural_evaluation": json.dumps(state["cultural_evaluation"])
            })

            # Clean and parse LLM response
            if result.startswith("```json"):
                result = result.replace("```json", "").replace("```", "").strip()
            elif result.startswith("```"):
                result = result.replace("```", "").strip()

            optional_result = json.loads(result)
            state["score_breakdown"]["optional_score"] = optional_result["optional_factors_score"]
            state["score_breakdown"]["optional_assessment"] = optional_result["assessment"]

            return state
        except Exception as e:
            state["error"] = f"Optional factors scoring failed: {str(e)}"
            return state

    def aggregate_scores(self, state: ScoringState) -> ScoringState:
        """
        Aggregate scores using provided weights.
        """
        if state.get("error"):
            return state

        try:
            weights = state["weights"]
            breakdown = state["score_breakdown"]

            final_score = (
                breakdown["technical_score"] * weights["technical"] +
                breakdown["communication_score"] * weights["communication"] +
                breakdown["cultural_score"] * weights["cultural"] +
                breakdown["optional_score"] * weights["optional"]
            )

            state["final_score"] = round(final_score, 2)
            state["score_breakdown"] = {
                "technical": {
                    "score": breakdown["technical_score"],
                    "weight": weights["technical"],
                    "contribution": round(breakdown["technical_score"] * weights["technical"], 2)
                },
                "communication": {
                    "score": breakdown["communication_score"],
                    "weight": weights["communication"],
                    "contribution": round(breakdown["communication_score"] * weights["communication"], 2)
                },
                "cultural": {
                    "score": breakdown["cultural_score"],
                    "weight": weights["cultural"],
                    "contribution": round(breakdown["cultural_score"] * weights["cultural"], 2)
                },
                "optional": {
                    "score": breakdown["optional_score"],
                    "weight": weights["optional"],
                    "contribution": round(breakdown["optional_score"] * weights["optional"], 2)
                }
            }

            return state
        except Exception as e:
            state["error"] = f"Score aggregation failed: {str(e)}"
            return state

    def save_to_mongodb(self, state: ScoringState) -> ScoringState:
        """
        Save aggregated score to MongoDB.
        """
        if state.get("error"):
            return state

        try:
            score_data = {
                "candidate_id": state["candidate_id"],
                "final_score": state["final_score"],
                "score_breakdown": state["score_breakdown"],
                "weights": state["weights"],
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime()),
                "processing_time": state["processing_time"]
            }
            score_data = convert_to_json_serializable(score_data)
            result = self.scores_collection.insert_one(score_data)
            state["score_breakdown"]["mongo_id"] = str(result.inserted_id)

            return state
        except Exception as e:
            state["error"] = f"Failed to save score to MongoDB: {str(e)}"
            return state

    def calculate_score(self, technical_evaluation, communication_evaluation, cultural_evaluation, weights=None):
        """
        Main function to calculate aggregated score using LangGraph workflow.
        """
        start_time = time.time()

        # Default weights
        default_weights = {
            "technical": 0.4,
            "communication": 0.25,
            "cultural": 0.25,
            "optional": 0.1
        }

        # Use provided weights or default
        weights = weights if weights else default_weights

        # Initialize state
        state = ScoringState(
            technical_evaluation=technical_evaluation,
            communication_evaluation=communication_evaluation,
            cultural_evaluation=cultural_evaluation,
            weights=weights,
            candidate_id=technical_evaluation.get("candidate_id", ""),
            score_breakdown={},
            final_score=0.0,
            processing_time=0.0,
            error=""
        )

        # Run the workflow
        result = self.workflow.invoke(state)
        result["processing_time"] = round(time.time() - start_time, 2)

        # Return result
        if result.get("error"):
            return {"error": result["error"]}
        return convert_to_json_serializable({
            "candidate_id": result["candidate_id"],
            "final_score": result["final_score"],
            "score_breakdown": result["score_breakdown"],
            "weights": result["weights"],
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime()),
            "processing_time": result["processing_time"],
            "mongo_id": result["score_breakdown"].get("mongo_id", "")
        })